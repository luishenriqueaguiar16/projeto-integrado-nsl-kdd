import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def main():
    print("=== Gerando Documento Word (.docx) do Relatório de Modelagem ===")
    
    doc = docx.Document()
    
    # Configurar margens ABNT (3cm superior/esquerda, 2cm inferior/direita)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)    # 3 cm
        section.left_margin = Inches(1.18)   # 3 cm
        section.bottom_margin = Inches(0.78) # 2 cm
        section.right_margin = Inches(0.78)  # 2 cm

    # Título Principal
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Relatório de Modelagem e Experimentação - NSL-KDD")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("🎓 Defesa Acadêmica Técnica (Módulo 3)")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(89, 89, 89)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()

    # Ler o conteúdo do relatório markdown
    md_path = os.path.join("projeto", "relatorio_modelagem.md")
    if not os.path.exists(md_path):
        print(f"Erro: Arquivo {md_path} não encontrado.")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    img_path = os.path.join("notebooks", "m3_modelagem", "imagens", "resultado_modelo_final.png")


    for line in lines:
        line_str = line.strip()
        
        if line_str.startswith("# Relatório") or line_str.startswith("### 🎓 Defesa"):
            continue
        elif line_str.startswith("## "):
            h = doc.add_heading(level=1)
            run = h.add_run(line_str.replace("## ", ""))
            run.font.name = "Calibri"
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif line_str.startswith("### "):
            h = doc.add_heading(level=2)
            run = h.add_run(line_str.replace("### ", ""))
            run.font.name = "Calibri"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182)
        elif line_str.startswith("#### "):
            h = doc.add_heading(level=3)
            run = h.add_run(line_str.replace("#### ", ""))
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
        elif line_str.startswith("![Resultado"):
            if os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(img_path, width=Inches(6.0))
                
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption_p.add_run("Figura 1: Matriz de Confusão e Curva ROC do Modelo Final no Conjunto de Teste")
                cap_run.font.size = Pt(9)
                cap_run.font.italic = True
                cap_run.font.color.rgb = RGBColor(89, 89, 89)
        elif line_str.startswith("* ") or line_str.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text = line_str[2:].replace("**", "").replace("*", "")
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. "):
            p = doc.add_paragraph()
            text = line_str.replace("**", "").replace("*", "")
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif line_str == "---" or not line_str:
            continue
        elif line_str.startswith("```"):
            continue
        else:
            p = doc.add_paragraph()
            text = line_str.replace("**", "").replace("*", "")
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    output_path = os.path.join("projeto", "relatorio_modelagem.docx")
    doc.save(output_path)
    print(f"[OK] Documento Word gerado com sucesso em: {output_path}")

    # Gerar v2 se existir
    md_v2_path = os.path.join("projeto", "relatorio_modelagem_v2.md")
    if os.path.exists(md_v2_path):
        doc_v2 = docx.Document()
        for section in doc_v2.sections:
            section.top_margin = Inches(1.18)
            section.left_margin = Inches(1.18)
            section.bottom_margin = Inches(0.78)
            section.right_margin = Inches(0.78)

        title_p = doc_v2.add_paragraph()
        title_run = title_p.add_run("Relatório de Modelagem e Experimentação - NSL-KDD (Versão Simplificada)")
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        with open(md_v2_path, "r", encoding="utf-8") as f:
            lines_v2 = f.readlines()

        for line in lines_v2:
            line_str = line.strip()
            if line_str.startswith("# Relatório") or line_str.startswith("### 🎓 Defesa"):
                continue
            elif line_str.startswith("## "):
                h = doc_v2.add_heading(level=1)
                run = h.add_run(line_str.replace("## ", ""))
                run.font.name = "Calibri"
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
            elif line_str.startswith("### "):
                h = doc_v2.add_heading(level=2)
                run = h.add_run(line_str.replace("### ", ""))
                run.font.name = "Calibri"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 117, 182)
            elif line_str.startswith("![Resultado"):
                if os.path.exists(img_path):
                    img_p = doc_v2.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_run = img_p.add_run()
                    img_run.add_picture(img_path, width=Inches(6.0))
                    
                    caption_p = doc_v2.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = caption_p.add_run("Figura 1: Matriz de Confusão e Curva ROC do Modelo Final no Conjunto de Teste")
                    cap_run.font.size = Pt(9)
                    cap_run.font.italic = True
                    cap_run.font.color.rgb = RGBColor(89, 89, 89)
            elif line_str.startswith("* ") or line_str.startswith("- "):
                p = doc_v2.add_paragraph(style='List Bullet')
                text = line_str[2:].replace("**", "").replace("*", "").replace("`", "")
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. ") or line_str.startswith("4. "):
                p = doc_v2.add_paragraph()
                text = line_str.replace("**", "").replace("*", "").replace("`", "")
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            elif line_str == "---" or not line_str:
                continue
            else:
                p = doc_v2.add_paragraph()
                text = line_str.replace("**", "").replace("*", "").replace("`", "")
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        output_v2_path = os.path.join("projeto", "relatorio_modelagem_v2.docx")
        doc_v2.save(output_v2_path)
        print(f"[OK] Documento Word v2 gerado com sucesso em: {output_v2_path}")

if __name__ == "__main__":
    main()
